import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Mm, Pt
import matplotlib.pyplot as plt
from io import BytesIO


from config import COMPANY_NAME


def create_report_word(figure: plt.figure,
                       data: dict,
                       report_filename: str ='report',
                       company_name: str = COMPANY_NAME,
                       ):
    doc = docx.Document()

    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    section.left_margin = Mm(25.4)
    section.right_margin = Mm(5.4)
    section.top_margin = Mm(5.4)
    section.bottom_margin = Mm(5.4)

    par0 = doc.add_paragraph(company_name)
    par0.add_run('\nИспытание саморегулирующегося нагревательного кабеля')
    par0.paragraph_format.space_before = Mm(1)
    par0.paragraph_format.space_after = Mm(1)
    par0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par1 = doc.add_paragraph(f'Дата - {data.get("s_TimePuskTest", "_")}')
    par1.add_run(f'\nФамилия испытателя - {data.get("s_FamilyTester", "_")}')
    par1.add_run(f'\nНаименование - {data.get("s_KableBrandTest", "_")}')
    par1.add_run(f'\n№ партии - {data.get("s_BatchNumberTest", "_")}')

    par1.paragraph_format.space_before = Mm(1)
    par1.paragraph_format.space_after = Mm(1)

    image = BytesIO()
    extent = figure.get_window_extent().transformed(figure.dpi_scale_trans.inverted())
    figure.savefig(image, format="jpg", bbox_inches=extent.expanded(0.85, 0.85))

    pic = doc.add_picture(image,
                          width=docx.shared.Cm(18),
                          height=docx.shared.Cm(15)
                          )
    #pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #pic.left_indent = docx.shared.Cm(105)
    #docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
    doc.save(report_filename)
